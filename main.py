
# Description:
# Author:
# Update Date:

import docx, copy, os,utils

# global variable
hidden_style = None
normal_style = None



#resource
text = utils.get_heyheyhey()


def run():
    #global
    global text


    # open docs return obj
    doc = docx.Document("./test.docx")

    normal_style = utils.get_style("./style.docx", "正文")
    hidden_style = utils.get_style("./style.docx", "嘿嘿嘿")

    for p0 in doc.paragraphs:
        p1 = copy.deepcopy(p0)
        p0.clear()
        for r in p1.runs:
            for t in r.text:
                p0.add_run(t, normal_style)
                p0.add_run(utils.insert(text), hidden_style)

    doc.save("done.docx")

def test():
    
    #global
    global text


    # open docs return obj
    doc = docx.Document("./test.docx")

    normal_style = utils.get_style("./style.docx", "正文")
    hidden_style = utils.get_style("./style.docx", "嘿嘿嘿")

    print(normal_style)
    print(hidden_style)

    # ot = []
    # it = "法律的概念其实很模糊法律所主张的正义依赖于判决而不是法律条文本身得益于消费者保护运动的"
    # for p in doc.paragraphs:
    #     for r in p.runs:
    #         for t in r.text:
    #             ot.append(t)

    # for p0 in doc.paragraphs:
    #     # p = p0.clear()
    #     for r in p0.runs:
    #         for t1 in r.text:
    #             p1.add_run(t1, normal_style)
    #             p1.add_run(insert(), hidden_style)
    # break
    # break

    for p0 in doc.paragraphs:
        p1 = copy.deepcopy(p0)
        p0.clear()
        for r in p1.runs:
            for t in r.text:
                p0.add_run(t, normal_style)
                # p0.add_run(insert(text), hidden_style)

    # print(type(p.text))
    # print(p.text[1])
    # if "嘿嘿" in p.text:
    # for r in p.runs:
    # print(r.style)
    # hidden_style = r.style
    # if "正文" in p.text:
    # for r in p.runs:
    # normal_style = r.style

    # for p in doc.paragraphs:
    #     p.add_run("正文插入", normal_style)
    #     p.add_run("隐藏插入", hidden_style)

    # doc.save("1.docx")





run()
# test()