# -*- coding: utf-8 -*-

import docx, copy, os


def get_heyheyhey():
    with open("./嘿嘿嘿.txt", encoding="utf-8") as f:
        s1 = f.read()
        for i in [
                " ", "，", "。", "\n", "\r", "（", "）", "[", "]", "？", "“", "”",
                "……", "——", "；", "《", "》", "、", "："
        ]:
            s1 = s1.replace(i, "")
        for i in range(10):
            s1 = s1.replace(str(i), "")
        # with open("./数据库.txt","a+") as f:
        #     f.write(s1)
        return s1

get_heyheyhey()

def get_style(style_path: str, style_name_in_file: str):
    """
    this is fun get the style obj
    @param style_path: like "./style.doc"
    @param style_name_in_file: like "正文"
    """
    doc = docx.Document(style_path)
    for p in doc.paragraphs:
        if style_name_in_file in p.text:
            for r in p.runs:
                # first result of finding
                return r.style


# count variable
insert_count = 0


def insert(text: str):
    global insert_count
    # text = "法律的概念其实很模糊法律所主张的正义依赖于判决而不是法律条文本身得益于消费者保护运动的"
    insert_count = insert_count + 1
    if (insert_count <= len(text)):
        return text[insert_count - 1]
    else:
        print("混淆字符用完了！")


# some object can't be serialized
def save_style_data(doc_path: str, style_list: list):
    """
    You must have a style docx which I can 
    1. Find one style with its text is its name
    @param: doc_path, like "./style.docx"
    @param: style_list, like ["正文","混淆字体"]
    """
    doc = docx.Document(doc_path)

    for style_name in style_list:
        for p in doc.paragraphs:
            if style_name in p.text:
                for r in p.runs:
                    with open("./style/" + style_name + ".pickle", "wb") as f:
                        # marshal.dump(r.style, f)
                        pass
                        break


def load_style_data(pickle_path: str):
    """
    I can search files in relative directory like "./style/*"
    pickle_path such as "style"
    """
    style_dict = {}
    for file_name in os.listdir("./" + pickle_path):
        with open("./" + pickle_path + "/" + file_name, "r") as f:
            # style_dict[file_name.split(".")[0]] = pickle.load(f)
            pass